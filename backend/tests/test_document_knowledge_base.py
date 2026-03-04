"""Tests for the Document Knowledge Base feature.

Covers:
  - document_parser.py    (file/URL → plain text)
  - document_store.py     (ChromaDB vector store)
  - documents_crud.py     (MongoDB CRUD)
  - document_indexing.py  (indexing pipeline)
  - kb_context.py         (context builder)
  - api/documents.py      (API endpoints)
"""
import base64
import io
import os
import sys
from datetime import datetime, timezone
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from bson import ObjectId
from fastapi import FastAPI
from fastapi.testclient import TestClient

# Ensure backend is on sys.path
TESTS_DIR = os.path.dirname(__file__)
BACKEND_DIR = os.path.abspath(os.path.join(TESTS_DIR, ".."))
if BACKEND_DIR not in sys.path:
    sys.path.insert(0, BACKEND_DIR)

import db.document_store as _ds_module


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture(autouse=True)
def _reset_document_store_singleton():
    """Reset the DocumentVectorStore singleton between tests."""
    _ds_module._document_store = None
    yield
    _ds_module._document_store = None


@pytest.fixture
def mock_documents_collection():
    """In-memory mongomock collection for documents."""
    import mongomock
    client = mongomock.MongoClient()
    return client["test_db"]["documents"]


# ---------------------------------------------------------------------------
# 1. document_parser.py
# ---------------------------------------------------------------------------

class TestDocumentParser:
    """Tests for document_parser.parse_file and parse_url."""

    def test_parse_txt(self):
        from services.document_parser import parse_file
        content = b"Hello, world!\nSecond line."
        text, ftype = parse_file("notes.txt", content)
        assert ftype == "txt"
        assert "Hello, world!" in text
        assert "Second line." in text

    def test_parse_md(self):
        from services.document_parser import parse_file
        content = b"# Heading\n\nSome **bold** text."
        text, ftype = parse_file("readme.md", content)
        assert ftype == "md"
        assert "Heading" in text

    def test_parse_docx(self):
        """Build a minimal DOCX in-memory and parse it."""
        from services.document_parser import parse_file
        from docx import Document as DocxDocument

        buf = io.BytesIO()
        doc = DocxDocument()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.save(buf)
        docx_bytes = buf.getvalue()

        text, ftype = parse_file("report.docx", docx_bytes)
        assert ftype == "docx"
        assert "First paragraph." in text
        assert "Second paragraph." in text

    def test_parse_pdf(self):
        """Build a minimal PDF in-memory via pdfplumber-compatible bytes."""
        from services.document_parser import parse_file
        # Use reportlab to create a simple PDF if available; otherwise skip
        pytest.importorskip("reportlab")
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.drawString(100, 750, "PDF test content for Zorbax protocol.")
        c.save()
        pdf_bytes = buf.getvalue()

        text, ftype = parse_file("spec.pdf", pdf_bytes)
        assert ftype == "pdf"
        assert "Zorbax" in text

    def test_parse_unsupported_extension(self):
        from services.document_parser import parse_file
        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_file("archive.zip", b"fake data")

    def test_parse_exceeds_max_chars(self):
        from services.document_parser import parse_file, MAX_DOCUMENT_CHARS
        large_content = ("A" * 201_000).encode()
        with pytest.raises(ValueError, match="too large"):
            parse_file("big.txt", large_content)

    def test_parse_url_success(self):
        from services.document_parser import parse_url
        # trafilatura is imported inside parse_url, so patch the module directly
        with patch("trafilatura.fetch_url", return_value="<html><body>Article.</body></html>"):
            with patch("trafilatura.extract", return_value="Article content here."):
                text = parse_url("https://example.com/article")
        assert text == "Article content here."

    def test_parse_url_fetch_failure(self):
        from services.document_parser import parse_url
        with patch("trafilatura.fetch_url", return_value=None):
            with pytest.raises(ValueError, match="Failed to fetch URL"):
                parse_url("https://bad-url.example")

    def test_parse_url_no_text_extracted(self):
        from services.document_parser import parse_url
        with patch("trafilatura.fetch_url", return_value="<html></html>"):
            with patch("trafilatura.extract", return_value=None):
                with pytest.raises(ValueError, match="No text content"):
                    parse_url("https://empty.example")


# ---------------------------------------------------------------------------
# 2. document_store.py
# ---------------------------------------------------------------------------

class TestDocumentVectorStore:
    """Tests for DocumentVectorStore."""

    def test_initialization(self):
        from db.document_store import DocumentVectorStore
        store = DocumentVectorStore(persist_directory=".chroma_docs_test")
        assert store.collection is not None
        assert store.client is not None

    def test_chunk_text_single_chunk(self):
        from db.document_store import DocumentVectorStore
        store = DocumentVectorStore(persist_directory=".chroma_docs_test")
        text = "Short text."
        chunks = store.chunk_text(text)
        assert len(chunks) >= 1
        assert all(isinstance(c, str) for c in chunks)

    def test_chunk_text_multiple_chunks(self):
        from db.document_store import DocumentVectorStore
        store = DocumentVectorStore(persist_directory=".chroma_docs_test")
        long_text = " ".join(["This is a test sentence."] * 200)
        chunks = store.chunk_text(long_text)
        assert len(chunks) > 1

    def test_add_and_delete_document(self):
        from db.document_store import DocumentVectorStore
        store = DocumentVectorStore(persist_directory=".chroma_docs_test")
        doc_id = "test_doc_001"
        chunks = ["First chunk of content.", "Second chunk of content."]
        embeddings = [[0.1] * 768, [0.2] * 768]
        metadata = {"title": "Test Doc", "source_type": "upload", "file_type": "txt"}

        success = store.add_document(doc_id, chunks, embeddings, metadata)
        assert success is True

        deleted = store.delete_document(doc_id)
        assert deleted is True

    def test_delete_nonexistent_document(self):
        from db.document_store import DocumentVectorStore
        store = DocumentVectorStore(persist_directory=".chroma_docs_test")
        # Should return False gracefully
        result = store.delete_document("nonexistent_doc_xyz")
        assert result is False

    def test_get_stats(self):
        from db.document_store import DocumentVectorStore
        store = DocumentVectorStore(persist_directory=".chroma_docs_test")
        stats = store.get_stats()
        assert "total_chunks" in stats
        assert "collection_name" in stats
        assert isinstance(stats["total_chunks"], int)

    def test_search_empty_collection_returns_empty(self):
        """Searching an empty collection should return [] without error."""
        from db.document_store import DocumentVectorStore
        store = DocumentVectorStore(persist_directory=".chroma_docs_test_empty")
        query_embedding = [0.1] * 768
        results = store.search_similar(query_embedding, limit=5, score_threshold=0.0)
        assert results == []

    def test_singleton(self):
        from db.document_store import get_document_store
        s1 = get_document_store(persist_directory=".chroma_docs_test")
        s2 = get_document_store(persist_directory=".chroma_docs_test")
        assert s1 is s2


# ---------------------------------------------------------------------------
# 3. documents_crud.py
# ---------------------------------------------------------------------------

class TestDocumentsCRUD:
    """Tests for MongoDB CRUD operations on documents collection."""

    def _make_doc(self, **overrides):
        base = {
            "user_id": "default",
            "title": "My Test Doc",
            "description": "A description",
            "source_type": "upload",
            "source_url": "",
            "file_type": "txt",
            "content": "Hello world content.",
            "char_count": 20,
            "indexed": False,
            "indexed_at": None,
            "created_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc),
        }
        base.update(overrides)
        return base

    def test_create_document(self, mock_documents_collection):
        from db.documents_crud import create_document
        with patch("db.documents_crud.get_documents_collection", return_value=mock_documents_collection):
            doc_id = create_document(self._make_doc())
        assert isinstance(doc_id, str)
        assert len(doc_id) == 24  # ObjectId hex string

    def test_get_document(self, mock_documents_collection):
        from db.documents_crud import create_document, get_document
        with patch("db.documents_crud.get_documents_collection", return_value=mock_documents_collection):
            doc_id = create_document(self._make_doc(title="Findable Doc"))
            doc = get_document(doc_id)
        assert doc is not None
        assert doc["title"] == "Findable Doc"

    def test_get_document_not_found(self, mock_documents_collection):
        from db.documents_crud import get_document
        with patch("db.documents_crud.get_documents_collection", return_value=mock_documents_collection):
            doc = get_document(str(ObjectId()))
        assert doc is None

    def test_list_documents(self, mock_documents_collection):
        from db.documents_crud import create_document, list_documents
        with patch("db.documents_crud.get_documents_collection", return_value=mock_documents_collection):
            create_document(self._make_doc(title="Doc A"))
            create_document(self._make_doc(title="Doc B"))
            docs = list_documents(user_id="default")
        assert len(docs) == 2

    def test_list_documents_filters_by_user(self, mock_documents_collection):
        from db.documents_crud import create_document, list_documents
        with patch("db.documents_crud.get_documents_collection", return_value=mock_documents_collection):
            create_document(self._make_doc(user_id="alice", title="Alice Doc"))
            create_document(self._make_doc(user_id="bob", title="Bob Doc"))
            alice_docs = list_documents(user_id="alice")
            bob_docs = list_documents(user_id="bob")
        assert len(alice_docs) == 1
        assert alice_docs[0]["title"] == "Alice Doc"
        assert len(bob_docs) == 1

    def test_update_document(self, mock_documents_collection):
        from db.documents_crud import create_document, get_document, update_document
        with patch("db.documents_crud.get_documents_collection", return_value=mock_documents_collection):
            doc_id = create_document(self._make_doc(title="Before"))
            update_document(doc_id, {"title": "After"})
            doc = get_document(doc_id)
        assert doc["title"] == "After"

    def test_delete_document(self, mock_documents_collection):
        from db.documents_crud import create_document, delete_document, get_document
        with patch("db.documents_crud.get_documents_collection", return_value=mock_documents_collection):
            doc_id = create_document(self._make_doc())
            deleted = delete_document(doc_id)
            assert deleted is True
            assert get_document(doc_id) is None

    def test_delete_document_not_found(self, mock_documents_collection):
        from db.documents_crud import delete_document
        with patch("db.documents_crud.get_documents_collection", return_value=mock_documents_collection):
            result = delete_document(str(ObjectId()))
        assert result is False

    def test_mark_document_indexed(self, mock_documents_collection):
        from db.documents_crud import create_document, get_document, mark_document_indexed
        with patch("db.documents_crud.get_documents_collection", return_value=mock_documents_collection):
            doc_id = create_document(self._make_doc(indexed=False))
            mark_document_indexed(doc_id)
            doc = get_document(doc_id)
        assert doc["indexed"] is True
        assert doc["indexed_at"] is not None

    def test_invalid_document_id_raises(self, mock_documents_collection):
        from db.documents_crud import get_document
        with patch("db.documents_crud.get_documents_collection", return_value=mock_documents_collection):
            with pytest.raises(ValueError, match="Invalid document id"):
                get_document("not-a-valid-id")


# ---------------------------------------------------------------------------
# 4. document_indexing.py
# ---------------------------------------------------------------------------

class TestDocumentIndexing:
    """Tests for the indexing pipeline."""

    @pytest.mark.asyncio
    async def test_index_document_success(self):
        from services.document_indexing import index_document
        mock_doc = {
            "_id": ObjectId(),
            "title": "Test Document",
            "content": "This is the document content for testing indexing.",
            "source_type": "upload",
            "file_type": "txt",
            "user_id": "default",
            "created_at": datetime.now(timezone.utc),
        }
        mock_store = MagicMock()
        mock_store.chunk_text.return_value = ["chunk one", "chunk two"]
        mock_store.add_document.return_value = True

        with patch("services.document_indexing.get_document", return_value=mock_doc):
            with patch("services.document_indexing.mark_document_indexed"):
                with patch("services.document_indexing.get_document_store", return_value=mock_store):
                    with patch("services.document_indexing.embed_batch", new_callable=AsyncMock) as mock_embed:
                        mock_embed.return_value = [[0.1] * 768, [0.2] * 768]
                        result = await index_document(str(mock_doc["_id"]))

        assert result is True
        mock_store.add_document.assert_called_once()

    @pytest.mark.asyncio
    async def test_index_document_not_found(self):
        from services.document_indexing import index_document
        with patch("services.document_indexing.get_document", return_value=None):
            result = await index_document(str(ObjectId()))
        assert result is False

    @pytest.mark.asyncio
    async def test_index_document_empty_content(self):
        from services.document_indexing import index_document
        mock_doc = {
            "_id": ObjectId(),
            "title": "Empty",
            "content": "",
            "source_type": "upload",
            "file_type": "txt",
        }
        with patch("services.document_indexing.get_document", return_value=mock_doc):
            result = await index_document(str(mock_doc["_id"]))
        assert result is False

    @pytest.mark.asyncio
    async def test_index_document_marks_indexed_on_success(self):
        from services.document_indexing import index_document
        mock_doc = {
            "_id": ObjectId(),
            "title": "Markable",
            "content": "Some content to index.",
            "source_type": "upload",
            "file_type": "txt",
            "user_id": "default",
            "created_at": datetime.now(timezone.utc),
        }
        mock_store = MagicMock()
        mock_store.chunk_text.return_value = ["single chunk"]
        mock_store.add_document.return_value = True

        mark_mock = MagicMock()
        with patch("services.document_indexing.get_document", return_value=mock_doc):
            with patch("services.document_indexing.mark_document_indexed", mark_mock):
                with patch("services.document_indexing.get_document_store", return_value=mock_store):
                    with patch("services.document_indexing.embed_batch", new_callable=AsyncMock) as mock_embed:
                        mock_embed.return_value = [[0.1] * 768]
                        await index_document(str(mock_doc["_id"]))

        mark_mock.assert_called_once_with(str(mock_doc["_id"]))

    @pytest.mark.asyncio
    async def test_index_document_does_not_mark_on_store_failure(self):
        from services.document_indexing import index_document
        mock_doc = {
            "_id": ObjectId(),
            "title": "Failing",
            "content": "Content that fails to store.",
            "source_type": "upload",
            "file_type": "txt",
            "user_id": "default",
            "created_at": datetime.now(timezone.utc),
        }
        mock_store = MagicMock()
        mock_store.chunk_text.return_value = ["single chunk"]
        mock_store.add_document.return_value = False  # Store fails

        mark_mock = MagicMock()
        with patch("services.document_indexing.get_document", return_value=mock_doc):
            with patch("services.document_indexing.mark_document_indexed", mark_mock):
                with patch("services.document_indexing.get_document_store", return_value=mock_store):
                    with patch("services.document_indexing.embed_batch", new_callable=AsyncMock) as mock_embed:
                        mock_embed.return_value = [[0.1] * 768]
                        result = await index_document(str(mock_doc["_id"]))

        assert result is False
        mark_mock.assert_not_called()


# ---------------------------------------------------------------------------
# 5. kb_context.py
# ---------------------------------------------------------------------------

class TestKBContext:
    """Tests for knowledge-base context building."""

    @pytest.mark.asyncio
    async def test_build_kb_context_returns_formatted_string(self):
        from services.kb_context import build_kb_context
        mock_embedding = [0.1] * 768
        mock_results = [
            {
                "id": "doc1_chunk_0",
                "doc_id": "doc1",
                "score": 0.88,
                "text": "The Zorbax protocol requires three daily check-ins.",
                "metadata": {"title": "Zorbax SOP", "source_type": "upload"},
            }
        ]
        with patch("services.kb_context.embed_text", return_value=mock_embedding):
            with patch("services.kb_context.get_document_store") as mock_get_store:
                mock_store = MagicMock()
                mock_get_store.return_value = mock_store
                mock_store.search_similar.return_value = mock_results
                ctx = await build_kb_context("Zorbax protocol", "default")

        assert "Knowledge Base" not in ctx  # header is added by chat_service, not here
        assert "Zorbax SOP" in ctx
        assert "88%" in ctx
        assert "Relevant documents" in ctx

    @pytest.mark.asyncio
    async def test_build_kb_context_empty_when_no_results(self):
        from services.kb_context import build_kb_context
        mock_embedding = [0.1] * 768
        with patch("services.kb_context.embed_text", return_value=mock_embedding):
            with patch("services.kb_context.get_document_store") as mock_get_store:
                mock_store = MagicMock()
                mock_get_store.return_value = mock_store
                mock_store.search_similar.return_value = []
                ctx = await build_kb_context("anything", "default")

        assert ctx == ""

    @pytest.mark.asyncio
    async def test_build_kb_context_short_query_returns_empty(self):
        from services.kb_context import build_kb_context
        ctx = await build_kb_context("hi", "default")
        assert ctx == ""

    @pytest.mark.asyncio
    async def test_build_kb_context_respects_char_limit(self):
        from services.kb_context import build_kb_context, MAX_KB_CONTEXT_CHARS
        mock_embedding = [0.1] * 768
        # 20 results with 400-char texts each — should be truncated
        mock_results = [
            {
                "id": f"doc{i}_chunk_0",
                "doc_id": f"doc{i}",
                "score": 0.8,
                "text": "B" * 400,
                "metadata": {"title": f"Doc {i}"},
            }
            for i in range(20)
        ]
        with patch("services.kb_context.embed_text", return_value=mock_embedding):
            with patch("services.kb_context.get_document_store") as mock_get_store:
                mock_store = MagicMock()
                mock_get_store.return_value = mock_store
                mock_store.search_similar.return_value = mock_results
                ctx = await build_kb_context("long query test", "default")

        assert len(ctx) <= MAX_KB_CONTEXT_CHARS

    @pytest.mark.asyncio
    async def test_build_kb_context_graceful_on_error(self):
        """Errors inside build_kb_context must not propagate."""
        from services.kb_context import build_kb_context
        with patch("services.kb_context.embed_text", side_effect=Exception("Ollama down")):
            ctx = await build_kb_context("test query", "default")
        assert ctx == ""


# ---------------------------------------------------------------------------
# 6. api/documents.py  (FastAPI endpoints via TestClient)
# ---------------------------------------------------------------------------

@pytest.fixture
def documents_test_client():
    """Minimal FastAPI app wired to the documents router."""
    from api.documents import router
    app = FastAPI()
    app.include_router(router)
    return TestClient(app, raise_server_exceptions=False)


def _b64(text: str) -> str:
    return base64.b64encode(text.encode()).decode()


class TestDocumentsAPI:
    """Tests for /api/documents endpoints."""

    # ---- /upload ----

    def test_upload_txt_success(self, documents_test_client):
        payload = {
            "filename": "notes.txt",
            "data_b64": _b64("Hello world content for testing."),
            "title": "Notes",
            "user_id": "default",
        }
        with patch("api.documents.parse_file", return_value=("Hello world content.", "txt")):
            with patch("api.documents.create_document", return_value=str(ObjectId())):
                with patch("api.documents.index_document", new_callable=AsyncMock, return_value=True):
                    resp = documents_test_client.post("/api/documents/upload", json=payload)

        assert resp.status_code == 200
        data = resp.json()
        assert data["success"] is True
        assert data["file_type"] == "txt"

    def test_upload_invalid_base64(self, documents_test_client):
        payload = {
            "filename": "bad.txt",
            "data_b64": "!!!not-base64!!!",
            "title": "Bad Upload",
        }
        resp = documents_test_client.post("/api/documents/upload", json=payload)
        assert resp.status_code == 400

    def test_upload_parse_error(self, documents_test_client):
        payload = {
            "filename": "bad.zip",
            "data_b64": _b64("fake zip"),
            "title": "Bad File",
        }
        with patch("api.documents.parse_file", side_effect=ValueError("Unsupported file type")):
            resp = documents_test_client.post("/api/documents/upload", json=payload)
        assert resp.status_code == 400
        assert "Unsupported" in resp.json()["detail"]

    # ---- /url ----

    def test_ingest_url_success(self, documents_test_client):
        with patch("api.documents.parse_url", return_value="Fetched web content here."):
            with patch("api.documents.create_document", return_value=str(ObjectId())):
                with patch("api.documents.index_document", new_callable=AsyncMock, return_value=True):
                    resp = documents_test_client.post(
                        "/api/documents/url",
                        json={"url": "https://example.com", "title": "Example"},
                    )
        assert resp.status_code == 200
        data = resp.json()
        assert data["success"] is True
        assert data["file_type"] == "url"

    def test_ingest_url_fetch_error(self, documents_test_client):
        with patch("api.documents.parse_url", side_effect=ValueError("Failed to fetch URL")):
            resp = documents_test_client.post(
                "/api/documents/url",
                json={"url": "https://bad.example", "title": "Bad"},
            )
        assert resp.status_code == 400

    # ---- GET /  (list) ----

    def test_list_documents(self, documents_test_client):
        mock_docs = [
            {
                "_id": ObjectId(),
                "user_id": "default",
                "title": "Doc 1",
                "file_type": "txt",
                "char_count": 100,
                "indexed": True,
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc),
            }
        ]
        with patch("api.documents.list_documents", return_value=mock_docs):
            resp = documents_test_client.get("/api/documents")
        assert resp.status_code == 200
        data = resp.json()
        assert data["count"] == 1
        assert data["documents"][0]["title"] == "Doc 1"

    def test_list_documents_excludes_content(self, documents_test_client):
        """The list endpoint must not include the 'content' field."""
        mock_docs = [
            {
                "_id": ObjectId(),
                "user_id": "default",
                "title": "Secret Doc",
                "file_type": "txt",
                "char_count": 5000,
                "content": "SECRET CONTENT SHOULD NOT APPEAR",
                "indexed": True,
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc),
            }
        ]
        with patch("api.documents.list_documents", return_value=mock_docs):
            resp = documents_test_client.get("/api/documents")
        assert "SECRET CONTENT" not in resp.text

    # ---- GET /{doc_id} ----

    def test_get_document_found(self, documents_test_client):
        doc_id = str(ObjectId())
        mock_doc = {
            "_id": ObjectId(doc_id),
            "title": "Found",
            "content": "Some content.",
            "file_type": "md",
            "created_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc),
        }
        with patch("api.documents.get_document", return_value=mock_doc):
            resp = documents_test_client.get(f"/api/documents/{doc_id}")
        assert resp.status_code == 200
        assert resp.json()["title"] == "Found"

    def test_get_document_not_found(self, documents_test_client):
        with patch("api.documents.get_document", return_value=None):
            resp = documents_test_client.get(f"/api/documents/{str(ObjectId())}")
        assert resp.status_code == 404

    # ---- DELETE /{doc_id} ----

    def test_delete_document_success(self, documents_test_client):
        doc_id = str(ObjectId())
        mock_doc = {"_id": ObjectId(doc_id), "title": "To delete"}
        mock_store = MagicMock()
        with patch("api.documents.get_document", return_value=mock_doc):
            with patch("api.documents.get_document_store", return_value=mock_store):
                with patch("api.documents.db_delete_document", return_value=True):
                    resp = documents_test_client.delete(f"/api/documents/{doc_id}")
        assert resp.status_code == 200
        assert resp.json()["success"] is True
        mock_store.delete_document.assert_called_once_with(doc_id)

    def test_delete_document_not_found(self, documents_test_client):
        with patch("api.documents.get_document", return_value=None):
            resp = documents_test_client.delete(f"/api/documents/{str(ObjectId())}")
        assert resp.status_code == 404

    # ---- GET /search ----

    def test_search_documents(self, documents_test_client):
        mock_results = [
            {
                "id": "doc1_chunk_0",
                "doc_id": "doc1",
                "score": 0.91,
                "text": "Matching content here.",
                "metadata": {"title": "Doc 1"},
            }
        ]
        mock_store = MagicMock()
        mock_store.search_similar.return_value = mock_results

        # embed_text is imported inside the search endpoint function, patch at its source
        with patch("api.documents.get_document_store", return_value=mock_store):
            with patch("services.embedder.embed_text", new_callable=AsyncMock, return_value=[0.1] * 768):
                resp = documents_test_client.get("/api/documents/search?q=matching")
        assert resp.status_code == 200

    def test_search_documents_missing_query(self, documents_test_client):
        resp = documents_test_client.get("/api/documents/search")
        # Missing required 'q' param → 422 Unprocessable Entity
        assert resp.status_code == 422

    # ---- GET /stats ----

    def test_get_stats(self, documents_test_client):
        mock_docs = [
            {"indexed": True, "char_count": 1000},
            {"indexed": False, "char_count": 500},
        ]
        mock_store = MagicMock()
        mock_store.get_stats.return_value = {"total_chunks": 15}
        with patch("api.documents.list_documents", return_value=mock_docs):
            with patch("api.documents.get_document_store", return_value=mock_store):
                resp = documents_test_client.get("/api/documents/stats")
        assert resp.status_code == 200
        data = resp.json()
        assert data["document_count"] == 2
        assert data["indexed_count"] == 1
        assert data["total_chars"] == 1500
        assert data["total_chunks"] == 15


# ---------------------------------------------------------------------------
# 7. Chat service integration (KB context injection)
# ---------------------------------------------------------------------------

class TestChatKBIntegration:
    """Verify chat_service wires in KB context."""

    def test_chat_service_imports_kb_context(self):
        """chat_service must be importable and reference build_kb_context."""
        import services.chat_service as cs
        import inspect
        src = inspect.getsource(cs)
        assert "build_kb_context" in src
        assert "=== Knowledge Base ===" in src

    def test_kb_section_included_in_sections_list(self):
        """The sections assembly must include kb_section."""
        import services.chat_service as cs
        import inspect
        src = inspect.getsource(cs)
        # Find the sections = [...] line
        for line in src.splitlines():
            if "sections = " in line and "guidelines_section" in line:
                assert "kb_section" in line, "kb_section missing from sections list"
                break
        else:
            pytest.fail("Could not find sections assembly line in chat_service.py")


# ---------------------------------------------------------------------------
# 8. KB Context for Tasks
# ---------------------------------------------------------------------------

class TestBuildTaskKBContext:
    """Tests for _build_task_kb_context() helper in api/tasks.py."""

    def _make_doc(self, doc_id, title, content, indexed=True):
        return {
            "_id": doc_id,
            "title": title,
            "content": content,
            "indexed": indexed,
            "file_type": "txt",
        }

    def test_returns_correct_kb_docs_and_context(self):
        """Valid indexed doc → kb_docs list + formatted kb_context string."""
        from api.tasks import _build_task_kb_context
        doc = self._make_doc("doc1", "SOP Document", "Follow steps A, B, C.")
        with patch("api.tasks._build_task_kb_context.__wrapped__" if hasattr(
            __import__("api.tasks", fromlist=["_build_task_kb_context"]),
            "__wrapped__"
        ) else "db.documents_crud.get_document", return_value=doc):
            with patch("db.documents_crud.get_document", return_value=doc):
                kb_docs, kb_context = _build_task_kb_context(["doc1"])
        assert len(kb_docs) == 1
        assert kb_docs[0] == {"id": "doc1", "title": "SOP Document"}
        assert "=== Knowledge Base ===" in kb_context
        assert "[SOP Document]" in kb_context
        assert "Follow steps A, B, C." in kb_context

    def test_skips_unindexed_doc(self):
        """Unindexed doc is silently skipped; returns empty."""
        from api.tasks import _build_task_kb_context
        doc = self._make_doc("doc1", "Draft", "some content", indexed=False)
        with patch("db.documents_crud.get_document", return_value=doc):
            kb_docs, kb_context = _build_task_kb_context(["doc1"])
        assert kb_docs == []
        assert kb_context == ""

    def test_skips_missing_doc(self):
        """get_document returning None is silently skipped."""
        from api.tasks import _build_task_kb_context
        with patch("db.documents_crud.get_document", return_value=None):
            kb_docs, kb_context = _build_task_kb_context(["missing_id"])
        assert kb_docs == []
        assert kb_context == ""

    def test_skips_on_exception(self):
        """get_document raising an exception is silently skipped."""
        from api.tasks import _build_task_kb_context
        with patch("db.documents_crud.get_document", side_effect=Exception("DB error")):
            kb_docs, kb_context = _build_task_kb_context(["bad_id"])
        assert kb_docs == []
        assert kb_context == ""

    def test_truncates_content_at_4000_chars(self):
        """Content longer than 4,000 chars is truncated to exactly 4,000."""
        from api.tasks import _build_task_kb_context, MAX_CHARS_PER_DOC
        long_content = "x" * 10_000
        doc = self._make_doc("doc1", "Big Doc", long_content)
        with patch("db.documents_crud.get_document", return_value=doc):
            kb_docs, kb_context = _build_task_kb_context(["doc1"])
        assert len(kb_docs) == 1
        excerpt = "x" * MAX_CHARS_PER_DOC
        assert excerpt in kb_context
        assert "x" * (MAX_CHARS_PER_DOC + 1) not in kb_context

    def test_enforces_max_2_docs(self):
        """Only the first 2 doc IDs are processed even if more are provided."""
        from api.tasks import _build_task_kb_context, MAX_KB_DOCS
        docs = {
            "d1": self._make_doc("d1", "Doc 1", "content 1"),
            "d2": self._make_doc("d2", "Doc 2", "content 2"),
            "d3": self._make_doc("d3", "Doc 3", "content 3"),
        }
        with patch("db.documents_crud.get_document", side_effect=lambda did: docs.get(did)):
            kb_docs, kb_context = _build_task_kb_context(["d1", "d2", "d3"])
        assert len(kb_docs) == MAX_KB_DOCS  # 2
        titles = [d["title"] for d in kb_docs]
        assert "Doc 1" in titles
        assert "Doc 2" in titles
        assert "Doc 3" not in titles

    def test_two_docs_both_appear_in_context(self):
        """Both docs appear in the combined kb_context string."""
        from api.tasks import _build_task_kb_context
        docs = {
            "d1": self._make_doc("d1", "Spec A", "spec content A"),
            "d2": self._make_doc("d2", "Spec B", "spec content B"),
        }
        with patch("db.documents_crud.get_document", side_effect=lambda did: docs.get(did)):
            kb_docs, kb_context = _build_task_kb_context(["d1", "d2"])
        assert len(kb_docs) == 2
        assert "[Spec A]" in kb_context
        assert "[Spec B]" in kb_context
        assert "spec content A" in kb_context
        assert "spec content B" in kb_context


class TestTaskPlannerKBContext:
    """Tests for kb_context parameter in task_planner functions."""

    def test_build_planning_prompt_includes_kb_text_when_provided(self):
        """build_planning_prompt includes the reference material block when kb_context is set."""
        from services.task_planner import build_planning_prompt, ALLOWED_TASK_TOOLS
        kb_context = "=== Knowledge Base ===\n[My SOP]\nDo step 1 then step 2."
        prompt = build_planning_prompt("Do a task", ALLOWED_TASK_TOOLS, None, kb_context=kb_context)
        assert "Reference material" in prompt
        assert "My SOP" in prompt
        assert "Do step 1 then step 2." in prompt

    def test_build_planning_prompt_omits_kb_text_when_empty(self):
        """build_planning_prompt does NOT include the reference material block when kb_context is empty."""
        from services.task_planner import build_planning_prompt, ALLOWED_TASK_TOOLS
        prompt = build_planning_prompt("Do a task", ALLOWED_TASK_TOOLS, None, kb_context="")
        assert "Reference material" not in prompt

    @pytest.mark.asyncio
    async def test_plan_task_passes_kb_context_to_prompt(self, monkeypatch):
        """plan_task passes kb_context through to build_planning_prompt."""
        import services.task_planner as planner
        captured = {}

        original = planner.build_planning_prompt
        def spy_prompt(goal, allowed_tools, budget, planner_hints=None, kb_context=""):
            captured["kb_context"] = kb_context
            return original(goal, allowed_tools, budget, planner_hints, kb_context)

        monkeypatch.setattr(planner, "build_planning_prompt", spy_prompt)

        async def fake_chat(messages, model, repeat_penalty=1.15):
            return '{"constraints":[],"resources":[],"steps":[{"id":"s1","title":"Search","instruction":"search","tool":"web_search","success_criteria":"found"}]}'

        monkeypatch.setattr(planner, "chat_with_provider", fake_chat)

        kb = "=== Knowledge Base ===\n[Guide]\nHere is the guide."
        await planner.plan_task("Do something", "llama3.1", None, kb_context=kb)
        assert captured["kb_context"] == kb


class TestTaskRunnerKBContext:
    """Tests for KB context injection in task_runner."""

    @pytest.mark.asyncio
    async def test_llm_step_injects_kb_context_as_second_message(self, monkeypatch):
        """When task has kb_context, LLM step messages include it after system message."""
        import services.task_runner as runner
        from db.tasks_crud import set_step_status, increment_usage

        kb_context = "=== Knowledge Base ===\n[Policy]\nNo refunds."
        task_doc = {
            "model_name": "llama3.1",
            "kb_context": kb_context,
            "planner_hints": None,
            "plan": {"steps": []},
        }
        captured_messages = {}

        async def fake_chat(messages, model):
            captured_messages["messages"] = messages
            return "LLM response"

        monkeypatch.setattr(runner, "chat_with_provider", fake_chat)
        monkeypatch.setattr(runner, "get_task", lambda tid: task_doc)
        monkeypatch.setattr(runner, "set_step_status", lambda *a, **kw: None)
        monkeypatch.setattr(runner, "increment_usage", lambda *a, **kw: None)
        monkeypatch.setattr(runner.progress_bus, "publish", AsyncMock())

        # Patch _verify_success to always pass
        async def fake_verify(*a, **kw):
            return True
        monkeypatch.setattr(runner, "_verify_success", fake_verify)

        step = {"tool": "llm.generate", "instruction": "Write something.", "success_criteria": "done", "params": {}, "timeout": 30, "max_retries": 0}
        await runner._execute_step("task1", 0, step)

        msgs = captured_messages.get("messages", [])
        assert msgs[0]["role"] == "system"
        assert msgs[1]["role"] == "user"
        assert msgs[1]["content"] == kb_context

    @pytest.mark.asyncio
    async def test_llm_step_skips_kb_injection_when_empty(self, monkeypatch):
        """When task has no kb_context, the message list starts with system + instruction only."""
        import services.task_runner as runner

        task_doc = {
            "model_name": "llama3.1",
            "kb_context": "",
            "planner_hints": None,
            "plan": {"steps": []},
        }
        captured_messages = {}

        async def fake_chat(messages, model):
            captured_messages["messages"] = messages
            return "LLM response"

        monkeypatch.setattr(runner, "chat_with_provider", fake_chat)
        monkeypatch.setattr(runner, "get_task", lambda tid: task_doc)
        monkeypatch.setattr(runner, "set_step_status", lambda *a, **kw: None)
        monkeypatch.setattr(runner, "increment_usage", lambda *a, **kw: None)
        monkeypatch.setattr(runner.progress_bus, "publish", AsyncMock())

        async def fake_verify(*a, **kw):
            return True
        monkeypatch.setattr(runner, "_verify_success", fake_verify)

        step = {"tool": "llm.generate", "instruction": "Write something.", "success_criteria": "done", "params": {}, "timeout": 30, "max_retries": 0}
        await runner._execute_step("task1", 0, step)

        msgs = captured_messages.get("messages", [])
        # System + prompt only (no KB message between them)
        kb_messages = [m for m in msgs if m.get("content", "").startswith("=== Knowledge Base ===")]
        assert kb_messages == []

    def test_task_runner_passes_kb_context_to_plan_task(self):
        """Source inspection: _run_task passes kb_context from task doc to plan_task."""
        import services.task_runner as runner
        import inspect
        src = inspect.getsource(runner._run_task)
        assert 'kb_context=doc.get("kb_context"' in src or "kb_context=" in src

    def test_task_runner_injects_kb_context_into_messages(self):
        """Source inspection: _execute_step reads kb_context from task_doc."""
        import services.task_runner as runner
        import inspect
        src = inspect.getsource(runner._execute_step)
        assert 'kb_context' in src
        assert 'messages.append' in src


class TestTaskAPIWithKBDocs:
    """Integration tests: task creation API stores kb_docs + kb_context."""

    def _make_tasks_api_client(self, mem):
        from fastapi import FastAPI
        from fastapi.testclient import TestClient
        import api.tasks as tasks_api
        app = FastAPI()
        tasks_api.create_task = lambda payload: mem.create_task(payload)
        tasks_api.get_task = lambda tid: mem.get_task(tid)
        tasks_api.list_tasks = lambda: mem.list_tasks()
        tasks_api.update_task = lambda tid, patch: mem.update_task(tid, patch)
        app.include_router(tasks_api.router)
        return TestClient(app)

    class _MemoryTasks:
        def __init__(self):
            self.docs = {}
        def create_task(self, doc):
            _id = str(len(self.docs) + 1)
            doc = dict(doc)
            doc["_id"] = _id
            self.docs[_id] = doc
            return _id
        def get_task(self, tid):
            return dict(self.docs.get(tid)) if tid in self.docs else None
        def list_tasks(self):
            return [dict(v) for v in self.docs.values()]
        def update_task(self, tid, patch):
            if tid in self.docs:
                self.docs[tid].update(patch)

    def test_task_created_without_kb_doc_ids_has_empty_kb_fields(self):
        """Task creation without kb_doc_ids → kb_docs=[] and kb_context=''."""
        mem = self._MemoryTasks()
        client = self._make_tasks_api_client(mem)
        r = client.post("/api/tasks", json={"goal": "do something", "dry_run": True})
        assert r.status_code == 200
        tid = r.json()["id"]
        doc = mem.get_task(tid)
        assert doc["kb_docs"] == []
        assert doc["kb_context"] == ""

    def test_task_created_with_kb_doc_ids_stores_context(self):
        """Task creation with valid kb_doc_ids → kb_docs and kb_context populated."""
        mem = self._MemoryTasks()
        client = self._make_tasks_api_client(mem)
        fake_doc = {
            "_id": "doc1",
            "title": "Research Paper",
            "content": "The findings show X.",
            "indexed": True,
        }
        with patch("api.tasks._build_task_kb_context", return_value=(
            [{"id": "doc1", "title": "Research Paper"}],
            "=== Knowledge Base ===\n[Research Paper]\nThe findings show X.",
        )):
            r = client.post("/api/tasks", json={
                "goal": "analyze research",
                "dry_run": True,
                "kb_doc_ids": ["doc1"],
            })
        assert r.status_code == 200
        tid = r.json()["id"]
        doc = mem.get_task(tid)
        assert len(doc["kb_docs"]) == 1
        assert doc["kb_docs"][0]["title"] == "Research Paper"
        assert "=== Knowledge Base ===" in doc["kb_context"]
        assert "Research Paper" in doc["kb_context"]

    def test_task_detail_includes_kb_docs_field(self):
        """GET /api/tasks/{id} returns kb_docs in the response."""
        mem = self._MemoryTasks()
        client = self._make_tasks_api_client(mem)
        with patch("api.tasks._build_task_kb_context", return_value=(
            [{"id": "doc1", "title": "My SOP"}],
            "=== Knowledge Base ===\n[My SOP]\ncontent here",
        )):
            r = client.post("/api/tasks", json={
                "goal": "follow the SOP",
                "dry_run": True,
                "kb_doc_ids": ["doc1"],
            })
        tid = r.json()["id"]
        r2 = client.get(f"/api/tasks/{tid}")
        assert r2.status_code == 200
        data = r2.json()
        assert "kb_docs" in data
        assert data["kb_docs"][0]["title"] == "My SOP"
