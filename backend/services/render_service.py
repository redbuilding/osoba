from __future__ import annotations

import io
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from core.config import BASE_DIR, get_logger

logger = get_logger("render_service")

try:
    import jinja2  # type: ignore
except Exception:
    jinja2 = None  # type: ignore

try:
    import markdown as md  # type: ignore
except Exception:
    md = None  # type: ignore

try:
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore
except Exception:
    Document = None  # type: ignore


@dataclass
class RenderCapabilities:
    html: bool
    docx: bool
    pdf: bool


def templates_dir() -> Path:
    return Path(BASE_DIR) / "templates"


def jinja_env() -> Optional["jinja2.Environment"]:
    if not jinja2:
        return None
    loader = jinja2.FileSystemLoader(str(templates_dir()))
    return jinja2.Environment(loader=loader, autoescape=True)


def render_markdown_to_html(markdown_text: str) -> str:
    if not md:
        raise RuntimeError("markdown library not installed")
    return md.markdown(markdown_text or "", extensions=["fenced_code", "tables", "toc", "sane_lists"])  # type: ignore


def render_html_template(template_path: str, context: dict) -> str:
    env = jinja_env()
    if not env:
        raise RuntimeError("Jinja2 not available")
    template = env.get_template(template_path)
    return template.render(**context)


def render_docx_from_markdown(markdown_text: str) -> bytes:
    if not Document:
        raise RuntimeError("python-docx not available")
    doc = Document()
    # Title: first line or default
    lines = (markdown_text or "").splitlines()
    title = lines[0].strip("# ") if lines else "Report"
    h = doc.add_heading(title, level=1)
    # Body: naive markdown flattening for now
    content = "\n".join(lines[1:]) if len(lines) > 1 else ""
    for para in content.split("\n\n"):
        doc.add_paragraph(para)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def pdf_enabled() -> bool:
    # PDF depends on playwright being installed and browsers set up externally (playwright install)
    try:
        import playwright  # type: ignore
        return True
    except Exception:
        return False


async def html_to_pdf_bytes(html: str) -> bytes:
    # Uses Playwright headless Chromium. Requires that the environment has browsers installed.
    try:
        from playwright.async_api import async_playwright  # type: ignore
    except Exception:
        raise RuntimeError("Playwright not available")

    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page()
        await page.set_content(html, wait_until="load")
        pdf_bytes = await page.pdf(format="Letter", print_background=True, margin={"top": "20mm", "right": "15mm", "bottom": "20mm", "left": "15mm"})
        await browser.close()
        return pdf_bytes


def get_capabilities() -> RenderCapabilities:
    return RenderCapabilities(
        html=bool(jinja2 and md),
        docx=bool(Document),
        pdf=bool(pdf_enabled()),
    )

